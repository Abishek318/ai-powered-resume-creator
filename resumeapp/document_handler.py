
from docx2pdf import convert
import base64
from docx import Document
import tempfile
import PyPDF2
import os
import uuid
# from .models import ResumeDocument
from django.core.files.base import ContentFile


def read_pdf(pdffile):
    text = ""
    # Iterate through each page and extract text
    reader=PyPDF2.PdfReader(pdffile.file)
    text=""
    for page in range(len(reader.pages)):
        page=reader.pages[page]
        text+=str(page.extract_text())
    return text

def word_bytes_to_pdf_base64(word_bytes,temp_id):
    # Create a unique filename
    unique_filename = str(uuid.uuid4())
    temp_dir = tempfile.gettempdir()
    docx_path = os.path.join(temp_dir, f"{unique_filename}.docx")
    pdf_path = os.path.join(temp_dir, f"{unique_filename}.pdf")
    try:
        # Save Word bytes to temporary file
        with open(docx_path, 'wb') as docx_file:
            docx_file.write(word_bytes)
        
        # Convert Word to PDF
        convert(docx_path, pdf_path)
        
        # Read the PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pdf_bytes = pdf_file.read()
        # try:
        #     pdf_document = ResumeDocument.objects.get(temp_id=temp_id)
        #     pdf_document.pdf.save('resume.pdf', ContentFile(pdf_bytes))
        # except ResumeDocument.DoesNotExist:
        #     pdf_document = ResumeDocument.objects.create(temp_id=temp_id)
        #     pdf_document.pdf.save('resume.pdf', ContentFile(pdf_bytes))
            # pdf_bytes=ResumeDocument.objects.get(temp_id=temp_id).pdf.read()
        # Encode PDF bytes to base64
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        return pdf_base64

    finally:
        # Clean up temporary files
        if os.path.exists(docx_path):
            os.remove(docx_path)
        if os.path.exists(pdf_path):
            os.remove(pdf_path)


def create_word_document(json_data):
    # Create a new Document
    doc = Document()

    # Add name as title
    doc.add_heading(json_data.get('name', 'Resume'), level=1)

    # Add contact information
    contact_info = json_data.get('contact_information', {})
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.add_run('Contact Information:').bold = True
    contact_paragraph.add_run(f"\nEmail: {contact_info.get('email', 'N/A')}")
    contact_paragraph.add_run(f"\nPhone: {contact_info.get('phone_number', 'N/A')}")
    contact_paragraph.add_run(f"\nAddress: {contact_info.get('address', 'N/A')}")

    # Add professional summary
    if 'professional_summary' in json_data:
        doc.add_heading('Professional Summary', level=2)
        doc.add_paragraph(json_data['professional_summary'])

    # Add education
    if 'education' in json_data:
        doc.add_heading('Education', level=2)
        for edu in json_data['education']:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', 'Degree')}, {edu.get('institution', 'Institution')}").bold = True
            p.add_run(f"\nGraduation Year: {edu.get('graduation_year', 'N/A')}")
            p.add_run(f"\nCGP: {edu.get('CGP', 'N/A')}")

    # Add work experience
    if 'work_experience' in json_data:
        doc.add_heading('Work Experience', level=2)
        for exp in json_data['work_experience']:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('job_title', 'Job Title')} at {exp.get('company', 'Company')}").bold = True
            p.add_run(f"\n{exp.get('start_date', 'Start Date')} - {exp.get('end_date', 'End Date')}")
            p.add_run(f"\nResponsibilities: {exp.get('responsibilities', 'N/A')}")

    # Add skills
    for skill_type in ['hard skills', 'soft skills']:
        if skill_type in json_data:
            doc.add_heading(skill_type.title(), level=2)
            doc.add_paragraph(', '.join(json_data[skill_type]))

    # Add certifications
    if 'certifications' in json_data:
        doc.add_heading('Certifications', level=2)
        for cert in json_data['certifications']:
            p = doc.add_paragraph()
            p.add_run(f"{cert.get('name', 'Certification Name')}").bold = True
            p.add_run(f" - {cert.get('institution', 'Institution')}, {cert.get('year', 'Year')}")

    # Add projects
    if 'projects' in json_data:
        doc.add_heading('Projects', level=2)
        for project in json_data['projects']:
            p = doc.add_paragraph()
            p.add_run(f"{project.get('name', 'Project Name')}").bold = True
            p.add_run(f"\nDescription: {project.get('description', 'N/A')}")
            p.add_run(f"\nTechnologies: {', '.join(project.get('technologies', []))}")

    # Add languages
    if 'languages' in json_data:
        doc.add_heading('Languages', level=2)
        doc.add_paragraph(', '.join(json_data['languages']))

    # Add awards
    if 'awards' in json_data:
        doc.add_heading('Awards', level=2)
        for award in json_data['awards']:
            p = doc.add_paragraph()
            p.add_run(f"{award.get('name', 'Award Name')}").bold = True
            p.add_run(f" - {award.get('institution', 'Institution')}, {award.get('year', 'Year')}")

    return doc



